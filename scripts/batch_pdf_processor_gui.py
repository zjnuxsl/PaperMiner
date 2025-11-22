#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
PaperMiner - 智能论文内容提取工具
功能：
1. 批量处理 input 文件夹中的所有 PDF
2. 自动提取文字、公式、图片、表格
3. 智能识别图片编号（Fig 1, Figure 2 等）
4. 智能章节提取（正则表达式 + LLM质量检查）
5. 生成规范的输出目录结构（extract文件夹）
"""

import tkinter as tk
from tkinter import ttk, messagebox
import subprocess
import threading
import sys
import io
import time
from pathlib import Path
import json
import re
import shutil
from typing import List

# 导入 LLM 辅助模块
try:
    from llm_helper import LLMHelper, load_prompt_template, save_sections
    LLM_AVAILABLE = True
except ImportError:
    LLM_AVAILABLE = False
    print("警告: 无法导入 LLM 模块，章节提取功能将不可用")

# 设置标准输出编码为 UTF-8
if sys.platform == 'win32':
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')
    sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding='utf-8', errors='replace')


class BatchPDFProcessorGUI:
    def __init__(self, root):
        self.root = root
        self.root.title("PaperMiner - 智能论文内容提取工具")
        self.root.geometry("960x720")

        # 设置基础路径
        self.base_path = Path(__file__).parent.parent
        self.input_path = self.base_path / "input"
        self.output_path = self.base_path / "output"
        self.raw_output_path = self.output_path / "raw"
        self.extract_output_path = self.output_path / "extract"

        # 处理状态
        self.is_processing = False
        self.current_pdf_index = 0
        self.total_pdfs = 0
        self.success_count = 0
        self.failed_count = 0
        # ������ɫ��������
        self.bg_color = '#F5F7FA'      # ���ڱ���
        self.card_bg = '#FFFFFF'       # �������/��ʾ����
        self.fg_color = '#333333'      # ���ı���ɫ
        self.accent_color = '#2D7BF4'  # ��ɫ����ɫ

        # 配置窗口
        self.root.minsize(960, 1200)
        self.root.grid_rowconfigure(0, weight=1)
        self.root.grid_columnconfigure(0, weight=1)

        # 配置样式
        self.setup_styles()

        self.create_widgets()
        self.check_input_folder()

    def setup_styles(self):
        """配置界面样式"""
        style = ttk.Style()

        # 设置主题
        try:
            style.theme_use('clam')
        except:
            pass

        # 配置颜色方案 - 简约风格
        bg_color = self.bg_color
        card_bg = self.card_bg
        fg_color = self.fg_color
        accent_color = self.accent_color

        # 配置各种控件样式
        style.configure('TFrame', background=bg_color)
        style.configure(
            'TLabel',
            background=bg_color,
            foreground=fg_color,
            font=('Microsoft YaHei UI', 10),
        )
        style.configure(
            'TLabelframe',
            background=card_bg,
            foreground='#555555',
            borderwidth=1,
            relief='solid',
        )
        style.configure(
            'TLabelframe.Label',
            background=card_bg,
            foreground='#555555',
            font=('Microsoft YaHei UI', 10, 'bold'),
        )

        # 按钮样式
        style.configure(
            'TButton',
            font=('Microsoft YaHei UI', 9),
            borderwidth=0,
            relief='flat',
            padding=(12, 6),
        )

        # 主要操作按钮样式
        style.configure(
            'Primary.TButton',
            font=('Microsoft YaHei UI', 10),
            background=accent_color,
            foreground='white',
            borderwidth=0,
            padding=(20, 10),
        )
        style.map(
            'Primary.TButton',
            background=[('active', '#1F66D1'), ('!disabled', accent_color)],
        )

        # 停止按钮样式
        style.configure(
            'Stop.TButton',
            font=('Microsoft YaHei UI', 10),
            background='#ffffff',
            foreground='#e74c3c',
            borderwidth=1,
            relief='solid',
            padding=(20, 10),
        )
        style.map(
            'Stop.TButton',
            background=[('active', '#e74c3c')],
            foreground=[('active', '#ffffff'), ('!disabled', '#e74c3c')],
        )

        # 复选框和单选按钮样式
        style.configure(
            'TCheckbutton',
            background=card_bg,
            foreground=fg_color,
            font=('Microsoft YaHei UI', 9),
        )
        style.map(
            'TCheckbutton',
            background=[('active', card_bg), ('!disabled', card_bg)],
            foreground=[('active', fg_color), ('!disabled', fg_color)],
        )

        style.configure(
            'TRadiobutton',
            background=card_bg,
            foreground=fg_color,
            font=('Microsoft YaHei UI', 9),
        )

        # 进度条样式
        style.configure(
            'TProgressbar',
            background=accent_color,
            troughcolor='#E5E9F2',
            borderwidth=0,
            thickness=18,
        )

        # 设置根窗口背景色
        self.root.configure(bg=bg_color)

    def create_styled_checkbutton(self, parent, text, variable):
        """创建统一样式的复选框（使用 tk.Checkbutton 以显示正确的勾选标记）"""
        return tk.Checkbutton(
            parent,
            text=text,
            variable=variable,
            bg=self.card_bg,
            fg=self.fg_color,
            font=('Microsoft YaHei UI', 9),
            activebackground=self.card_bg,
            activeforeground=self.fg_color,
            selectcolor='white',
            relief='flat',
            highlightthickness=0
        )

    def create_widgets(self):
        """创建界面组件"""
        main_frame = ttk.Frame(self.root, padding="20")
        main_frame.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))
        main_frame.grid_columnconfigure(0, weight=1)
        main_frame.grid_rowconfigure(4, weight=1)

        # 标题区域 - 简约风格
        title_frame = tk.Frame(main_frame, bg=self.accent_color, height=60)
        title_frame.grid(row=0, column=0, pady=(0, 20), sticky=(tk.W, tk.E))
        title_frame.grid_propagate(False)

        title_label = tk.Label(
            title_frame,
            text="📄 PaperMiner - 智能论文内容提取工具",
            font=('Microsoft YaHei UI', 16, 'bold'),
            bg=self.accent_color,
            fg='white'
        )
        title_label.pack(expand=True)
        title_label.config(text="PaperMiner - 智能论文内容提取工具")
        
        # 文件信息区域
        info_frame = ttk.LabelFrame(main_frame, text="📁 文件信息", padding="12")
        info_frame.grid(row=1, column=0, pady=(0, 12), sticky=(tk.W, tk.E))
        info_frame.configure(text="📄 文件信息")
        info_frame.grid_columnconfigure(1, weight=1)

        self.file_count_label = ttk.Label(
            info_frame,
            text="PDF 文件数量: 0",
            font=('Microsoft YaHei UI', 10)
        )
        self.file_count_label.grid(row=0, column=0, sticky=tk.W, pady=4)

        ttk.Button(
            info_frame,
            text="🔄 刷新",
            command=self.check_input_folder,
            width=10
        ).grid(row=0, column=1, sticky=tk.E, pady=4)

        ttk.Button(
            info_frame,
            text="📂 打开 input 文件夹",
            command=lambda: self.open_folder(self.input_path)
        ).grid(row=1, column=0, columnspan=2, sticky=(tk.W, tk.E), pady=4)
        
        # 处理选项区域
        options_frame = ttk.LabelFrame(main_frame, text="⚙️ 处理选项", padding="12")
        options_frame.grid(row=2, column=0, pady=(0, 12), sticky=(tk.W, tk.E))
        options_frame.configure(text="⚙ 处理选项")

        # 处理模式选择
        self.process_mode_var = tk.StringVar(value="full")

        mode_frame = ttk.Frame(options_frame)
        mode_frame.grid(row=0, column=0, columnspan=2, sticky=(tk.W, tk.E), pady=(0, 10))

        ttk.Label(
            mode_frame,
            text="处理模式:",
            font=('Microsoft YaHei UI', 9)
        ).pack(side=tk.LEFT, padx=(0, 10))

        ttk.Radiobutton(
            mode_frame,
            text="完整处理 (PDF → MinerU → 提取)",
            variable=self.process_mode_var,
            value="full",
            command=self.on_mode_change
        ).pack(side=tk.LEFT, padx=(0, 15))

        ttk.Radiobutton(
            mode_frame,
            text="仅提取 (从已有 raw 文件夹提取)",
            variable=self.process_mode_var,
            value="extract_only",
            command=self.on_mode_change
        ).pack(side=tk.LEFT)

        # 分隔线
        ttk.Separator(options_frame, orient='horizontal').grid(
            row=1, column=0, columnspan=2, sticky=(tk.W, tk.E), pady=10
        )

        self.extract_text_var = tk.BooleanVar(value=True)
        self.extract_formula_var = tk.BooleanVar(value=True)
        self.extract_figures_var = tk.BooleanVar(value=True)
        self.extract_tables_var = tk.BooleanVar(value=True)
        self.extract_sections_var = tk.BooleanVar(value=False)  # 默认不勾选
        self.use_gpu_var = tk.BooleanVar(value=True)

        # 使用统一样式的复选框（显示正确的勾选标记）
        self.create_styled_checkbutton(
            options_frame,
            "✏️ 提取文字 (Markdown)",
            self.extract_text_var
        ).grid(row=2, column=0, sticky=tk.W, pady=4)

        self.create_styled_checkbutton(
            options_frame,
            "🔢 提取公式 (LaTeX)",
            self.extract_formula_var
        ).grid(row=3, column=0, sticky=tk.W, pady=4)

        self.create_styled_checkbutton(
            options_frame,
            "🖼️提取图片(智能识别编号)",
            self.extract_figures_var
        ).grid(row=4, column=0, sticky=tk.W, pady=4)

        self.create_styled_checkbutton(
            options_frame,
            "📊 提取表格 (Excel)",
            self.extract_tables_var
        ).grid(row=5, column=0, sticky=tk.W, pady=4)

        self.create_styled_checkbutton(
            options_frame,
            "📑 提取论文章节 (正则表达式 + LLM)",
            self.extract_sections_var
        ).grid(row=6, column=0, sticky=tk.W, pady=4)

        # LLM 模型选择（已移除，固定使用 Deepseek）
        # llm_frame = ttk.Frame(options_frame)
        # llm_frame.grid(row=7, column=0, sticky=tk.W, padx=(30, 0), pady=2)
        #
        # ttk.Label(llm_frame, text="LLM 模型:").pack(side=tk.LEFT, padx=(0, 5))
        #
        self.llm_model_var = tk.StringVar(value="deepseek")
        # llm_combo = ttk.Combobox(
        #     llm_frame,
        #     textvariable=self.llm_model_var,
        #     values=["deepseek"],
        #     state="readonly",
        #     width=12
        # )
        # llm_combo.pack(side=tk.LEFT)

        ttk.Separator(options_frame, orient='horizontal').grid(
            row=8, column=0, sticky=(tk.W, tk.E), pady=10
        )

        self.gpu_checkbox = self.create_styled_checkbutton(
            options_frame,
            "⚡ 使用 GPU 加速 (推荐)",
            self.use_gpu_var
        )
        self.gpu_checkbox.grid(row=9, column=0, sticky=tk.W, pady=4)
        
        # 控制按钮区域
        control_frame = ttk.Frame(main_frame)
        control_frame.grid(row=3, column=0, pady=(0, 12), sticky=(tk.W, tk.E))
        control_frame.grid_columnconfigure(0, weight=1)
        control_frame.grid_columnconfigure(1, weight=1)

        self.start_button = ttk.Button(
            control_frame,
            text="▶  开始处理",
            command=self.start_processing,
            style='Primary.TButton'
        )
        self.start_button.grid(row=0, column=0, sticky=(tk.W, tk.E), padx=(0, 6))
        self.start_button.config(text="开始处理")

        self.stop_button = ttk.Button(
            control_frame,
            text="⏹  停止",
            command=self.stop_processing,
            state='disabled',
            style='Stop.TButton'
        )
        self.stop_button.grid(row=0, column=1, sticky=(tk.W, tk.E), padx=(6, 0))
        self.stop_button.config(text="停止")

        # 进度区域
        progress_frame = ttk.LabelFrame(main_frame, text="📊 处理进度", padding="12")
        progress_frame.grid(row=4, column=0, pady=(0, 12), sticky=(tk.W, tk.E, tk.N, tk.S))
        progress_frame.configure(text="📊 处理进度")
        progress_frame.grid_columnconfigure(0, weight=1)
        progress_frame.grid_rowconfigure(3, weight=1)

        self.status_label = ttk.Label(
            progress_frame,
            text="就绪",
            font=('Microsoft YaHei UI', 10),
            foreground='#27ae60'
        )
        self.status_label.grid(row=0, column=0, sticky=tk.W, pady=(0, 8))
        self.status_label.config(text="准备就绪")

        self.progress_var = tk.DoubleVar()
        self.progress_bar = ttk.Progressbar(
            progress_frame,
            variable=self.progress_var,
            maximum=100,
            mode='determinate'
        )
        self.progress_bar.grid(row=1, column=0, sticky=(tk.W, tk.E), pady=(0, 8))

        self.progress_text = ttk.Label(
            progress_frame,
            text="当前 0 / 总计 0",
            font=('Microsoft YaHei UI', 9)
        )
        self.progress_text.grid(row=2, column=0, sticky=tk.W, pady=(0, 8))

        # 统计信息区域 - 简约风格
        stats_frame = ttk.Frame(progress_frame)
        stats_frame.grid(row=3, column=0, sticky=(tk.W, tk.E), pady=(4, 0))
        stats_frame.grid_columnconfigure(0, weight=1)
        stats_frame.grid_columnconfigure(1, weight=1)
        stats_frame.grid_columnconfigure(2, weight=1)

        # 成功统计
        success_frame = tk.Frame(stats_frame, bg='#e8f5e9', relief='solid', borderwidth=1)
        success_frame.grid(row=0, column=0, sticky=(tk.W, tk.E), padx=(0, 4))
        tk.Label(
            success_frame,
            text="成功",
            font=('Microsoft YaHei UI', 9),
            bg='#e8f5e9',
            fg='#666'
        ).pack(pady=(6, 2))
        self.success_label = tk.Label(
            success_frame,
            text="0",
            font=('Microsoft YaHei UI', 14, 'bold'),
            bg='#e8f5e9',
            fg='#27ae60'
        )
        self.success_label.pack(pady=(2, 6))

        # 失败统计
        failed_frame = tk.Frame(stats_frame, bg='#ffebee', relief='solid', borderwidth=1)
        failed_frame.grid(row=0, column=1, sticky=(tk.W, tk.E), padx=4)
        tk.Label(
            failed_frame,
            text="失败",
            font=('Microsoft YaHei UI', 9),
            bg='#ffebee',
            fg='#666'
        ).pack(pady=(6, 2))
        self.failed_label = tk.Label(
            failed_frame,
            text="0",
            font=('Microsoft YaHei UI', 14, 'bold'),
            bg='#ffebee',
            fg='#e74c3c'
        )
        self.failed_label.pack(pady=(2, 6))

        # 总计统计
        total_frame = tk.Frame(stats_frame, bg='#f5f5f5', relief='solid', borderwidth=1)
        total_frame.grid(row=0, column=2, sticky=(tk.W, tk.E), padx=(4, 0))
        tk.Label(
            total_frame,
            text="总计",
            font=('Microsoft YaHei UI', 9),
            bg='#f5f5f5',
            fg='#666'
        ).pack(pady=(6, 2))
        self.total_label = tk.Label(
            total_frame,
            text="0",
            font=('Microsoft YaHei UI', 14, 'bold'),
            bg='#f5f5f5',
            fg='#333'
        )
        self.total_label.pack(pady=(2, 6))

        # 输出目录按钮区域
        output_frame = ttk.LabelFrame(main_frame, text="📂 输出目录", padding="12")
        output_frame.grid(row=5, column=0, pady=(0, 12), sticky=(tk.W, tk.E))
        output_frame.configure(text="📂 输出目录")
        output_frame.grid_columnconfigure(0, weight=1)
        output_frame.grid_columnconfigure(1, weight=1)

        ttk.Button(
            output_frame,
            text="📁 打开原始输出 (raw)",
            command=lambda: self.open_folder(self.raw_output_path)
        ).grid(row=0, column=0, sticky=(tk.W, tk.E), padx=(0, 6), pady=4)

        ttk.Button(
            output_frame,
            text="📁 打开提取结果 (extract)",
            command=lambda: self.open_folder(self.extract_output_path)
        ).grid(row=0, column=1, sticky=(tk.W, tk.E), padx=(6, 0), pady=4)

        # 底部信息
        footer_frame = ttk.Frame(main_frame)
        footer_frame.grid(row=6, column=0, sticky=(tk.W, tk.E), pady=(8, 0))

        ttk.Label(
            footer_frame,
            text="邮箱:sl-xiao@zjnu.cn",
            font=('Microsoft YaHei UI', 8),
            foreground='#999'
        ).pack(side=tk.LEFT)

    def check_input_folder(self):
        """检查 input 文件夹中的 PDF 文件"""
        if not self.input_path.exists():
            self.input_path.mkdir(parents=True, exist_ok=True)
            self.file_count_label.config(text="PDF 文件数量: 0")
            self.log("⚠️  input 文件夹为空，请添加 PDF 文件")
            return

        pdf_files = list(self.input_path.glob("*.pdf"))
        count = len(pdf_files)
        self.file_count_label.config(text=f"PDF 文件数量: {count}")

        if count > 0:
            self.log(f"✓ 找到 {count} 个 PDF 文件")
            for pdf in pdf_files:
                self.log(f"  - {pdf.name}")
        else:
            self.log("⚠️  input 文件夹为空，请添加 PDF 文件")

    def log(self, message: str):
        """添加日志消息（仅输出到控制台）"""
        print(message, flush=True)  # 输出到控制台

    def open_folder(self, folder_path: Path):
        """打开文件夹"""
        if folder_path.exists():
            import os
            os.startfile(str(folder_path))
        else:
            messagebox.showwarning("警告", f"目录不存在：{folder_path}")

    def on_mode_change(self):
        """处理模式切换"""
        mode = self.process_mode_var.get()

        if mode == "extract_only":
            # 仅提取模式：禁用 GPU 选项
            self.gpu_checkbox.config(state='disabled')
            self.use_gpu_var.set(False)
        else:
            # 完整处理模式：启用 GPU 选项
            self.gpu_checkbox.config(state='normal')
            self.use_gpu_var.set(True)

    def start_processing(self):
        """开始处理"""
        mode = self.process_mode_var.get()

        # 检查是否至少选择了一项提取内容
        if not any([
            self.extract_text_var.get(),
            self.extract_formula_var.get(),
            self.extract_figures_var.get(),
            self.extract_tables_var.get(),
            self.extract_sections_var.get()
        ]):
            messagebox.showwarning(
                "未选择提取项",
                "请至少选择一项提取内容！"
            )
            return

        if mode == "full":
            # 完整处理模式：需要 PDF 文件
            pdf_files = list(self.input_path.glob("*.pdf"))
            if not pdf_files:
                messagebox.showwarning(
                    "没有文件",
                    "input 文件夹中没有 PDF 文件！\n\n请先添加 PDF 文件。"
                )
                return

            # 构建提取内容描述
            extract_items = []
            if self.extract_text_var.get():
                extract_items.append("文字")
            if self.extract_formula_var.get():
                extract_items.append("公式")
            if self.extract_figures_var.get():
                extract_items.append("图片")
            if self.extract_tables_var.get():
                extract_items.append("表格")
            if self.extract_sections_var.get():
                extract_items.append("论文章节")

            extract_desc = "、".join(extract_items)

            # 确认开始处理
            if not messagebox.askyesno(
                "确认处理",
                f"将完整处理 {len(pdf_files)} 个 PDF 文件。\n\n"
                f"提取项目：{extract_desc}\n\n"
                f"这将运行 MinerU 并提取内容。\n\n是否继续？"
            ):
                return

            items_to_process = pdf_files

        else:  # extract_only
            # 仅提取模式：检查 raw 文件夹
            if not self.raw_output_path.exists():
                messagebox.showwarning(
                    "没有 raw 文件夹",
                    f"raw 文件夹不存在：{self.raw_output_path}\n\n"
                    f"请先运行完整处理模式生成 raw 文件夹。"
                )
                return

            # 查找所有 raw 子文件夹
            raw_folders = [d for d in self.raw_output_path.iterdir() if d.is_dir()]
            if not raw_folders:
                messagebox.showwarning(
                    "没有数据",
                    f"raw 文件夹中没有数据！\n\n"
                    f"请先运行完整处理模式。"
                )
                return

            # 构建提取内容描述
            extract_items = []
            if self.extract_text_var.get():
                extract_items.append("文字")
            if self.extract_formula_var.get():
                extract_items.append("公式")
            if self.extract_figures_var.get():
                extract_items.append("图片")
            if self.extract_tables_var.get():
                extract_items.append("表格")
            if self.extract_sections_var.get():
                extract_items.append("论文章节")

            extract_desc = "、".join(extract_items)

            # 确认开始提取
            if not messagebox.askyesno(
                "确认提取",
                f"将从 {len(raw_folders)} 个 raw 文件夹中提取内容。\n\n"
                f"提取项目：{extract_desc}\n\n"
                f"这将跳过 MinerU 处理，直接从已有的 raw 文件夹提取。\n\n是否继续？"
            ):
                return

            items_to_process = raw_folders

        # 更新界面状态
        self.is_processing = True
        self.start_button.config(state='disabled')
        self.stop_button.config(state='normal')

        # 重置计数器和统计显示
        self.current_pdf_index = 0
        self.total_pdfs = len(items_to_process)
        self.success_count = 0
        self.failed_count = 0

        # 更新统计显示
        self.update_stats()

        # 在后台线程中处理
        if mode == "full":
            threading.Thread(
                target=self.process_pdfs,
                args=(items_to_process,),
                daemon=True
            ).start()
        else:
            threading.Thread(
                target=self.extract_from_raw,
                args=(items_to_process,),
                daemon=True
            ).start()

    def stop_processing(self):
        """停止处理"""
        self.is_processing = False
        self.log("\n⚠️  用户请求停止处理...")

    def process_pdfs(self, pdf_files: List[Path]):
        """处理所有 PDF 文件"""
        try:
            self.log("=" * 60)
            self.log("PaperMiner - 批量 PDF 处理开始")
            self.log("=" * 60)
            self.log(f"总文件数: {len(pdf_files)}")
            self.log(f"输出目录: {self.output_path}")
            self.log("")

            # 创建输出目录
            self.raw_output_path.mkdir(parents=True, exist_ok=True)
            self.extract_output_path.mkdir(parents=True, exist_ok=True)

            # 检查 GPU 状态
            self.check_gpu_status()

            # 处理每个 PDF
            for i, pdf_file in enumerate(pdf_files):
                if not self.is_processing:
                    self.log("\n❌ 处理已停止")
                    break

                self.current_pdf_index = i + 1
                self.update_progress()

                self.log("\n" + "=" * 60)
                self.log(f"[{i+1}/{len(pdf_files)}] 处理: {pdf_file.name}")
                self.log("=" * 60)

                # 步骤 1: 使用 MinerU 处理 PDF
                success = self.run_mineru(pdf_file)

                if success:
                    # 步骤 2: 提取和整理结果
                    self.extract_and_organize(pdf_file.stem)
                    self.success_count += 1
                    self.log(f"✅ 完成: {pdf_file.name}")
                else:
                    self.failed_count += 1
                    self.log(f"❌ 失败: {pdf_file.name}")

                # 更新统计显示
                self.root.after(0, lambda: self.update_stats())

            # 处理完成
            self.log("\n" + "=" * 60)
            self.log("处理完成!")
            self.log("=" * 60)
            self.log(f"成功: {self.success_count} 个")
            self.log(f"失败: {self.failed_count} 个")
            self.log(f"总计: {len(pdf_files)} 个")

            self.root.after(0, lambda: self.processing_complete())

        except Exception as e:
            self.log(f"\n❌ 处理过程中发生错误: {str(e)}")
            import traceback
            self.log(traceback.format_exc())
            self.root.after(0, lambda: self.processing_complete())

    def update_progress(self):
        """更新进度"""
        if self.total_pdfs > 0:
            progress = (self.current_pdf_index / self.total_pdfs) * 100
            self.progress_var.set(progress)
            self.progress_text.config(
                text=f"{self.current_pdf_index} / {self.total_pdfs}"
            )
            self.progress_text.config(
                text=f"当前 {self.current_pdf_index} / 总计 {self.total_pdfs}"
            )
            self.status_label.config(
                text=f"正在处理... ({self.current_pdf_index}/{self.total_pdfs})",
                foreground=self.accent_color
            )

    def update_stats(self):
        """更新统计信息显示"""
        self.success_label.config(text=str(self.success_count))
        self.failed_label.config(text=str(self.failed_count))
        self.total_label.config(text=str(self.total_pdfs))

    def check_gpu_status(self):
        """检查 GPU 状态"""
        try:
            self.log("=== GPU 诊断 ===")

            result = subprocess.run([
                sys.executable, "-c",
                "import torch; "
                "print('PyTorch:', torch.__version__); "
                "print('CUDA:', torch.cuda.is_available()); "
                "print('GPU:', torch.cuda.get_device_name(0) if torch.cuda.is_available() else 'N/A')"
            ], capture_output=True, text=True, encoding='utf-8', errors='replace', timeout=30)

            if result.returncode == 0:
                for line in result.stdout.strip().split('\n'):
                    self.log(line)

                # 检查 CUDA 是否可用
                cuda_check = subprocess.run([
                    sys.executable, "-c",
                    "import torch; exit(0 if torch.cuda.is_available() else 1)"
                ], capture_output=True, timeout=10)

                if cuda_check.returncode != 0:
                    self.log("⚠️  CUDA 不可用，将使用 CPU 模式")
                    self.use_gpu_var.set(False)
                else:
                    self.log("✅ GPU 加速已启用")
            else:
                self.log("⚠️  GPU 检查失败")

            self.log("=" * 60)
            self.log("")

        except Exception as e:
            self.log(f"⚠️  GPU 状态检查失败: {str(e)}")
            self.log("")

    def run_mineru(self, pdf_file: Path) -> bool:
        """运行 MinerU 处理 PDF（实时显示输出）"""
        try:
            self.log("步骤 1: 使用 MinerU 处理 PDF...")

            # 构建命令
            device = "cuda" if self.use_gpu_var.get() else "cpu"

            # 使用 mineru 命令
            cmd = ['mineru', '-p', str(pdf_file), '-o', str(self.raw_output_path), '-d', device]

            self.log(f"命令: {' '.join(cmd)}")
            self.log("正在处理，请稍候...")
            self.log("")

            # 使用 Popen 实时显示输出
            process = subprocess.Popen(
                cmd,
                stdout=subprocess.PIPE,
                stderr=subprocess.STDOUT,  # 合并 stderr 到 stdout
                text=True,
                encoding='utf-8',
                errors='replace',
                cwd=str(self.base_path),
                bufsize=1,  # 行缓冲
                universal_newlines=True
            )

            # 实时读取输出
            last_log_time = 0
            line_count = 0

            for line in process.stdout:
                line = line.strip()
                if not line:
                    continue

                line_count += 1
                current_time = time.time()

                # 显示所有包含关键信息的行
                if any(keyword in line.lower() for keyword in [
                    'processing', 'page', 'error', 'warning', 'success',
                    'complete', '完成', 'info', 'progress', '处理',
                    'layout', 'ocr', 'model', 'cuda', 'gpu'
                ]):
                    self.log(f"  {line}")
                # 每5秒显示一次进度提示（即使没有关键词）
                elif current_time - last_log_time > 5:
                    self.log(f"  处理中... (已读取 {line_count} 行输出)")
                    last_log_time = current_time

                # 强制刷新GUI
                self.root.update_idletasks()

            # 等待进程完成
            return_code = process.wait(timeout=3600)  # 60分钟超时

            if return_code == 0:
                self.log("")
                self.log("✓ MinerU 处理完成")
                return True
            else:
                self.log("")
                self.log(f"❌ MinerU 处理失败 (返回码: {return_code})")
                return False

        except subprocess.TimeoutExpired:
            self.log("")
            self.log("❌ 处理超时 (60 分钟)")
            if 'process' in locals():
                process.kill()
            return False
        except Exception as e:
            self.log("")
            self.log(f"❌ MinerU 处理失败: {str(e)}")
            import traceback
            self.log(traceback.format_exc())
            return False

    def extract_from_raw(self, raw_folders: List[Path]):
        """从已有的 raw 文件夹中提取内容"""
        try:
            self.log("=" * 60)
            self.log("仅提取模式 - 从 raw 文件夹提取内容")
            self.log("=" * 60)
            self.log(f"总文件夹数: {len(raw_folders)}")
            self.log(f"输出目录: {self.extract_output_path}")
            self.log("=" * 60)
            self.log("")

            for i, raw_folder in enumerate(raw_folders, 1):
                if not self.is_processing:
                    self.log("\n⚠️  处理已停止")
                    break

                self.current_pdf_index = i
                pdf_name = raw_folder.name

                self.log("=" * 60)
                self.log(f"[{i}/{len(raw_folders)}] 提取: {pdf_name}")
                self.log("=" * 60)

                # 更新进度
                self.root.after(0, lambda: self.update_progress())
                self.root.after(0, lambda name=pdf_name: self.status_label.config(
                    text=f"正在提取 {name}...",
                    foreground=self.accent_color
                ))

                try:
                    # 提取内容（使用 PDF 名称作为参数）
                    self.extract_and_organize(pdf_name)

                    self.success_count += 1
                    self.log(f"✅ 完成: {pdf_name}")
                    self.log("")

                except Exception as e:
                    self.failed_count += 1
                    self.log(f"❌ 失败: {pdf_name}")
                    self.log(f"   错误: {str(e)}")
                    import traceback
                    self.log(traceback.format_exc())
                    self.log("")

                # 更新统计显示
                self.root.after(0, lambda: self.update_stats())

            # 处理完成
            self.log("=" * 60)
            self.log("提取完成!")
            self.log("=" * 60)
            self.log(f"成功: {self.success_count} 个")
            self.log(f"失败: {self.failed_count} 个")
            self.log(f"总计: {len(raw_folders)} 个")

            self.processing_complete()

        except Exception as e:
            self.log(f"❌ 提取过程出错: {str(e)}")
            self.processing_complete()

    def extract_and_organize(self, pdf_name: str):
        """提取和整理处理结果到extract文件夹"""
        try:
            self.log("步骤 2: 提取和整理结果...")

            # 查找 MinerU 的输出目录
            raw_pdf_dir = self.raw_output_path / pdf_name / "auto"
            if not raw_pdf_dir.exists():
                self.log(f"⚠️  未找到输出目录: {raw_pdf_dir}")
                return

            # 创建提取目录（在extract下为每个PDF创建子文件夹）
            extract_pdf_dir = self.extract_output_path / pdf_name
            extract_pdf_dir.mkdir(parents=True, exist_ok=True)

            # 提取文字 (Markdown) - 保存到extract/pdf_name/pdf_name.md
            if self.extract_text_var.get():
                self.extract_text(raw_pdf_dir, extract_pdf_dir, pdf_name)

            # 提取公式 - 保存到extract/pdf_name/Formula/
            if self.extract_formula_var.get():
                self.extract_formulas(raw_pdf_dir, extract_pdf_dir, pdf_name)

            # 提取图片 - 保存到extract/pdf_name/Figure/
            if self.extract_figures_var.get():
                self.extract_figures(raw_pdf_dir, extract_pdf_dir, pdf_name)

            # 提取表格 - 保存到extract/pdf_name/Tables/
            if self.extract_tables_var.get():
                self.extract_tables(raw_pdf_dir, extract_pdf_dir, pdf_name)

            # 提取论文章节 (使用 LLM) - 保存到extract/pdf_name/Sections/
            if self.extract_sections_var.get():
                self.extract_sections_with_llm(raw_pdf_dir, extract_pdf_dir, pdf_name)

            # 创建 Word 文件夹（按原文顺序排列图片和表格） - 保存到extract/pdf_name/Word/
            if self.extract_figures_var.get() or self.extract_tables_var.get():
                self.create_word_folder(raw_pdf_dir, extract_pdf_dir, pdf_name)

            self.log("✓ 提取和整理完成")

        except Exception as e:
            self.log(f"❌ 提取失败: {str(e)}")
            import traceback
            self.log(traceback.format_exc())

    def processing_complete(self):
        """处理完成"""
        self.is_processing = False
        self.start_button.config(state='normal')
        self.stop_button.config(state='disabled')

        # 更新状态标签
        if self.failed_count == 0:
            self.status_label.config(
                text="处理完成",
                foreground='#27ae60'
            )
        elif self.success_count == 0:
            self.status_label.config(
                text="处理失败",
                foreground='#e74c3c'
            )
        else:
            self.status_label.config(
                text="处理完成（部分失败）",
                foreground='#f39c12'
            )

        # 更新统计显示
        self.update_stats()

        # 显示完成消息
        messagebox.showinfo(
            "处理完成",
            f"批量处理完成！\n\n"
            f"成功: {self.success_count} 个\n"
            f"失败: {self.failed_count} 个\n"
            f"总计: {self.total_pdfs} 个\n\n"
            f"结果已保存到 output/extract 文件夹"
        )

    def extract_text(self, raw_dir: Path, extract_dir: Path, pdf_name: str):
        """提取文字 (Markdown)，并修复图片引用路径"""
        try:
            self.log("  - 提取文字...")

            md_file = raw_dir / f"{pdf_name}.md"
            if not md_file.exists():
                self.log("    ⚠️  未找到 Markdown 文件")
                return

            # 读取 Markdown 文件
            with open(md_file, 'r', encoding='utf-8') as f:
                content = f.read()

            # 修复图片引用路径
            # 将 images/xxx.jpg 替换为 Figure/xxx.jpg
            # 同时需要处理表格图片和公式图片

            # 读取 content_list.json 来区分图片类型
            content_list_file = raw_dir / f"{pdf_name}_content_list.json"
            image_mapping = {}  # 原始路径 -> 新路径

            if content_list_file.exists():
                with open(content_list_file, 'r', encoding='utf-8') as f:
                    content_list = json.load(f)

                # 构建图片映射
                fig_index = 1
                table_index = 1
                formula_index = 1

                for item in content_list:
                    if 'img_path' in item:
                        img_path = item['img_path']
                        item_type = item.get('type')

                        if item_type == 'image':
                            # 普通图片 -> Figure/
                            # 尝试从标题中提取图片编号
                            caption = ' '.join(item.get('img_caption', []))
                            fig_match = re.search(r'(Fig\.?|Figure)\s*(\d+)', caption, re.IGNORECASE)
                            if fig_match:
                                new_path = f"Figure/Fig.{fig_match.group(2)}.jpg"
                            else:
                                new_path = f"Figure/image_{fig_index}.jpg"
                            fig_index += 1
                        elif item_type == 'table':
                            # 表格图片 -> Tables/
                            caption = ' '.join(item.get('table_caption', []))
                            table_match = re.search(r'Table\s*(\d+)', caption, re.IGNORECASE)
                            if table_match:
                                new_path = f"Tables/Table_{table_match.group(1)}.jpg"
                            else:
                                new_path = f"Tables/Table_{table_index}.jpg"
                            table_index += 1
                        elif item_type == 'equation':
                            # 公式图片 -> Formula/
                            new_path = f"Formula/formula_{formula_index}.jpg"
                            formula_index += 1
                        else:
                            continue

                        image_mapping[img_path] = new_path

            # 替换图片路径
            for old_path, new_path in image_mapping.items():
                # 处理 Markdown 图片引用格式: ![xxx](images/xxx.jpg)
                content = content.replace(f"]({old_path})", f"]({new_path})")

            # 保存修改后的 Markdown 文件
            output_md = extract_dir / f"{pdf_name}.md"
            with open(output_md, 'w', encoding='utf-8') as f:
                f.write(content)

            self.log(f"    ✓ 已保存: {output_md.name}")
            self.log(f"    ✓ 修复了 {len(image_mapping)} 个图片引用")

        except Exception as e:
            self.log(f"    ❌ 文字提取失败: {str(e)}")
            import traceback
            self.log(traceback.format_exc())

    def extract_formulas(self, raw_dir: Path, extract_dir: Path, pdf_name: str):
        """提取公式（从 content_list.json 和 Markdown）"""
        try:
            self.log("  - 提取公式...")

            # 读取 content_list.json
            content_list_file = raw_dir / f"{pdf_name}_content_list.json"
            formula_images = []

            if content_list_file.exists():
                with open(content_list_file, 'r', encoding='utf-8') as f:
                    content_list = json.load(f)

                # 提取所有公式图片
                for item in content_list:
                    if item.get('type') == 'equation':
                        if 'img_path' in item:
                            formula_images.append({
                                'img_path': item['img_path'],
                                'latex': item.get('latex_text', '')
                            })

            # 同时从 Markdown 文件中提取文本公式
            md_file = raw_dir / f"{pdf_name}.md"
            text_formulas = []
            if md_file.exists():
                with open(md_file, 'r', encoding='utf-8') as f:
                    content = f.read()
                # 提取公式块 ($$...$$)
                text_formulas = re.findall(r'\$\$(.*?)\$\$', content, re.DOTALL)

            if not formula_images and not text_formulas:
                self.log("    ⚠️  未找到公式")
                return

            # 创建公式文件夹
            formula_dir = extract_dir / "Formula"
            formula_dir.mkdir(exist_ok=True)

            # 保存公式图片
            if formula_images:
                for i, formula_item in enumerate(formula_images, 1):
                    img_path = raw_dir / formula_item['img_path']
                    if img_path.exists():
                        dest_path = formula_dir / f"formula_{i}{img_path.suffix}"
                        shutil.copy2(img_path, dest_path)

            # 保存文本公式到 Markdown 文件
            if text_formulas:
                formula_md = formula_dir / f"{pdf_name}_formula.md"
                with open(formula_md, 'w', encoding='utf-8') as f:
                    f.write(f"# {pdf_name} - 公式\n\n")
                    for i, formula in enumerate(text_formulas, 1):
                        f.write(f"## 公式 {i}\n\n")
                        f.write(f"$$\n{formula.strip()}\n$$\n\n")

            total_formulas = len(formula_images) + len(text_formulas)
            self.log(f"    ✓ 提取 {total_formulas} 个公式 (图片: {len(formula_images)}, 文本: {len(text_formulas)})")

        except Exception as e:
            self.log(f"    ❌ 公式提取失败: {str(e)}")
            import traceback
            self.log(traceback.format_exc())

    def extract_figures(self, raw_dir: Path, extract_dir: Path, pdf_name: str):
        """提取图片并智能识别编号（排除公式和表格图片）"""
        try:
            self.log("  - 提取图片...")

            # 读取 content_list.json 以区分图片类型
            content_list_file = raw_dir / f"{pdf_name}_content_list.json"
            figure_images = []
            excluded_images = set()  # 公式和表格图片

            if content_list_file.exists():
                with open(content_list_file, 'r', encoding='utf-8') as f:
                    content_list = json.load(f)

                # 提取图片（figures）
                for item in content_list:
                    item_type = item.get('type')
                    if item_type == 'image':
                        # 真正的图片
                        if 'img_path' in item:
                            figure_images.append({
                                'img_path': item['img_path'],
                                'caption': item.get('img_caption', [])
                            })
                    elif item_type in ['equation', 'table']:
                        # 公式和表格图片，需要排除
                        if 'img_path' in item:
                            excluded_images.add(item['img_path'])

            if not figure_images:
                self.log("    ⚠️  未找到图片")
                return

            self.log(f"    找到 {len(figure_images)} 张图片（排除 {len(excluded_images)} 张公式/表格图片）")

            # 创建图片文件夹
            figure_dir = extract_dir / "Figure"
            figure_dir.mkdir(exist_ok=True)

            # 复制并重命名图片，同时创建映射文件
            copied_count = 0
            image_mapping = {}  # 原始路径 -> 新文件名的映射

            for i, fig_item in enumerate(figure_images, 1):
                try:
                    img_path = raw_dir / fig_item['img_path']
                    if not img_path.exists():
                        self.log(f"    ⚠️  图片不存在: {img_path}")
                        continue

                    # 尝试从 caption 中提取图片编号
                    new_name = None
                    if fig_item['caption']:
                        caption_text = ' '.join(fig_item['caption'])
                        # 匹配图片标题格式
                        caption_match = re.match(
                            r'^(Fig\.?|Figure|图|Scheme|示意图)\s*(\d+)',
                            caption_text,
                            re.IGNORECASE
                        )
                        if caption_match:
                            fig_num = caption_match.group(2)
                            new_name = f"Fig.{fig_num}{img_path.suffix}"

                    # 如果没有找到编号，使用默认命名
                    if not new_name:
                        new_name = f"image_{i}{img_path.suffix}"

                    # 复制图片
                    output_img = figure_dir / new_name
                    shutil.copy2(img_path, output_img)
                    copied_count += 1

                    # 保存映射关系
                    image_mapping[fig_item['img_path']] = new_name

                except Exception as e:
                    self.log(f"    ⚠️  复制图片失败: {str(e)}")

            # 保存图片映射到 JSON 文件
            mapping_file = figure_dir / "image_mapping.json"
            with open(mapping_file, 'w', encoding='utf-8') as f:
                json.dump(image_mapping, f, ensure_ascii=False, indent=2)

            self.log(f"    ✓ 提取 {copied_count} 张图片")

        except Exception as e:
            self.log(f"    ❌ 图片提取失败: {str(e)}")
            import traceback
            self.log(traceback.format_exc())

    def extract_tables(self, raw_dir: Path, extract_dir: Path, pdf_name: str):
        """提取表格（从 content_list.json，保存为 Excel 和 JPG）"""
        try:
            self.log("  - 提取表格...")

            # 读取 content_list.json
            content_list_file = raw_dir / f"{pdf_name}_content_list.json"
            if not content_list_file.exists():
                self.log("    ⚠️  未找到 content_list.json 文件")
                return

            with open(content_list_file, 'r', encoding='utf-8') as f:
                content_list = json.load(f)

            # 提取所有表格
            tables = []
            for item in content_list:
                if item.get('type') == 'table':
                    tables.append(item)

            if not tables:
                self.log("    ⚠️  未找到表格")
                return

            self.log(f"    找到 {len(tables)} 个表格")

            # 创建 Tables 文件夹（与 Figure、Formula 同级）
            tables_dir = extract_dir / "Tables"
            tables_dir.mkdir(exist_ok=True)

            # 保存表格
            for i, table_item in enumerate(tables, 1):
                # 获取表格标题和编号
                table_caption = ""
                table_number = str(i)

                if 'table_caption' in table_item and table_item['table_caption']:
                    caption_text = ' '.join(table_item['table_caption'])
                    table_caption = caption_text

                    # 尝试从标题中提取表格编号
                    caption_match = re.search(r'Table\s*(\d+)', caption_text, re.IGNORECASE)
                    if caption_match:
                        table_number = caption_match.group(1)

                # 1. 保存表格图片（始终保存）
                if 'img_path' in table_item:
                    img_path = raw_dir / table_item['img_path']
                    if img_path.exists():
                        img_dest = tables_dir / f"Table_{table_number}{img_path.suffix}"
                        shutil.copy2(img_path, img_dest)
                        self.log(f"    ✓ 保存表格图片 {table_number}: {img_dest.name}")

                # 2. 保存表格为 Excel（Sheet 名称为文件名，表头显示完整标题）
                try:
                    import pandas as pd
                    from bs4 import BeautifulSoup

                    if 'table_body' in table_item:
                        # 解析 HTML 表格
                        soup = BeautifulSoup(table_item['table_body'], 'html.parser')
                        table = soup.find('table')

                        if table:
                            # 提取表格数据
                            rows = []
                            for tr in table.find_all('tr'):
                                row = [td.get_text(strip=True) for td in tr.find_all(['td', 'th'])]
                                rows.append(row)

                            # 创建 DataFrame
                            if rows:
                                # 方法1：尝试使用第一行作为表头
                                try:
                                    df = pd.DataFrame(rows[1:], columns=rows[0] if len(rows) > 1 else None)
                                except ValueError as e:
                                    # 列数不匹配（通常是因为 rowspan/colspan）
                                    # 方法2：找到列数最多的行作为参考
                                    max_cols = max(len(row) for row in rows)
                                    
                                    # 尝试使用列数最多的行作为参考，不使用表头
                                    self.log(f"    ⓘ 表格 {table_number} 结构复杂（rowspan/colspan），使用备用方案")
                                    
                                    # 统一列数：不足的行用空字符串填充
                                    normalized_rows = []
                                    for row in rows:
                                        if len(row) < max_cols:
                                            row = row + [''] * (max_cols - len(row))
                                        normalized_rows.append(row[:max_cols])  # 截取到最大列数
                                    
                                    # 不使用表头，使用默认列名（Column 1, Column 2, ...）
                                    df = pd.DataFrame(normalized_rows)

                                # Sheet 名称使用文件名（Table_1）
                                sheet_name = f"Table_{table_number}"

                                # 保存为 Excel
                                excel_file = tables_dir / f"Table_{table_number}.xlsx"

                                # 使用 ExcelWriter 来添加表格标题
                                with pd.ExcelWriter(excel_file, engine='openpyxl') as writer:
                                    df.to_excel(writer, index=False, sheet_name=sheet_name)

                                    # 获取工作表
                                    worksheet = writer.sheets[sheet_name]

                                    # 在第一行插入表格标题（如果有）
                                    if table_caption:
                                        worksheet.insert_rows(1)
                                        worksheet['A1'] = table_caption
                                        # 合并第一行的单元格
                                        from openpyxl.styles import Font, Alignment
                                        worksheet.merge_cells(start_row=1, start_column=1, end_row=1, end_column=len(df.columns))
                                        worksheet['A1'].font = Font(bold=True, size=12)
                                        worksheet['A1'].alignment = Alignment(horizontal='center', vertical='center')

                                self.log(f"    ✓ 保存表格 Excel {table_number}: {excel_file.name}")

                except ImportError:
                    self.log("    ⚠️  需要安装 pandas 和 openpyxl 来生成 Excel 文件")
                    self.log("    ⚠️  运行: pip install pandas openpyxl beautifulsoup4")
                except Exception as e:
                    self.log(f"    ⚠️  表格 {table_number} Excel 生成失败: {str(e)}")

            self.log(f"    ✓ 提取 {len(tables)} 个表格")

        except Exception as e:
            self.log(f"    ❌ 表格提取失败: {str(e)}")
            import traceback
            self.log(traceback.format_exc())

    def extract_sections_with_llm(self, raw_dir: Path, extract_dir: Path, pdf_name: str):
        """提取论文章节（优先使用正则表达式，失败时使用 LLM）"""
        try:
            self.log("  - 提取论文章节...")

            # 检查 LLM 模块是否可用
            if not LLM_AVAILABLE:
                self.log("    ❌ LLM 模块不可用，请检查 llm_helper.py 是否存在")
                return

            # 读取 Markdown 文件
            md_file = extract_dir / f"{pdf_name}.md"
            if not md_file.exists():
                # 尝试从 raw 目录读取
                md_file = raw_dir / f"{pdf_name}.md"

            if not md_file.exists():
                self.log("    ❌ 未找到 Markdown 文件，请先提取文字")
                return

            self.log("    读取 Markdown 文件...")
            with open(md_file, 'r', encoding='utf-8') as f:
                markdown_content = f.read()

            # 检查文件大小（避免超过 API 限制）
            content_length = len(markdown_content)
            self.log(f"    文档长度: {content_length} 字符")

            if content_length > 100000:
                self.log("    ⚠️  文档过长，可能超过 API 限制，尝试截取前 100000 字符")
                markdown_content = markdown_content[:100000]

            # 提示词模板将在需要时加载（根据是否有缺失章节选择不同模板）
            prompt_template = None

            # 初始化 LLM
            model_name = self.llm_model_var.get()
            self.log(f"    使用模型: {model_name.upper()}")

            try:
                llm = LLMHelper(model_name=model_name)
            except ValueError as e:
                self.log(f"    ❌ LLM 初始化失败: {str(e)}")
                self.log("    💡 请检查 .env 文件中的 API Key 配置")
                return

            # 优先使用正则表达式提取章节（同时获取未识别的标题）
            self.log("    🔄 使用正则表达式提取章节...")
            sections, unrecognized_headers = llm.extract_sections_fallback(markdown_content, return_unrecognized=True)

            # 质量检查：判断是否需要使用 LLM
            need_llm = False
            llm_reason = []

            if not sections:
                # 情况1: 正则表达式完全失败
                need_llm = True
                llm_reason.append("正则表达式未能提取任何章节")
            else:
                # 情况2: 检查关键章节是否缺失
                critical_sections = ['Abstract', 'Introduction', 'Methods', 'Results & Discussion', 'Conclusion']
                missing_critical = [s for s in critical_sections if s not in sections]
                
                if missing_critical:
                    need_llm = True
                    llm_reason.append(f"缺少关键章节: {', '.join(missing_critical)}")
                
                # 情况3: 检查章节内容是否过短（可能截断）
                short_sections = []
                for name, content in sections.items():
                    if len(content.strip()) < 100:  # 少于100字符
                        short_sections.append(name)
                
                if short_sections:
                    need_llm = True
                    llm_reason.append(f"章节内容过短: {', '.join(short_sections)}")
                
                # 情况4: 章节数量异常
                if len(sections) < 2:
                    need_llm = True
                    llm_reason.append(f"章节数量过少 ({len(sections)}个)")
                elif len(sections) > 8:
                    # 章节过多可能是误识别，但不强制使用LLM
                    self.log(f"    ⓘ 提取到较多章节 ({len(sections)}个)，可能包含子章节")

            # 优先尝试使用 LLM 对未识别的标题进行分类（轻量级方案）
            if unrecognized_headers and sections:
                self.log(f"    🔍 发现 {len(unrecognized_headers)} 个未识别的章节标题，尝试使用 LLM 分类...")
                try:
                    classification = llm.classify_section_titles(unrecognized_headers)

                    if classification:
                        # 根据分类结果，从 markdown 中提取对应章节内容
                        lines = markdown_content.split('\n')

                        # 按章节类型分组（处理多个标题映射到同一章节类型的情况）
                        sections_by_type = {}
                        for header_text, section_type in classification.items():
                            if section_type not in sections_by_type:
                                sections_by_type[section_type] = []
                            sections_by_type[section_type].append(header_text)

                        # 提取每个章节类型的内容
                        for section_type, headers in sections_by_type.items():
                            # 检查该章节是否已经被正则提取过
                            already_exists = section_type in sections
                            if already_exists:
                                self.log(f"       ⓘ {section_type} 已存在（正则提取），将合并 LLM 分类的内容")

                            # 提取所有匹配该类型的标题的内容，并合并
                            combined_content = []

                            # 找到所有标题在文档中的位置
                            header_positions = []
                            for header_text in headers:
                                for line_idx, line in enumerate(lines):
                                    if line.strip() == header_text:
                                        header_positions.append((line_idx, header_text))
                                        break

                            # 按位置排序
                            header_positions.sort(key=lambda x: x[0])

                            # 提取内容：从第一个标题到最后一个标题之后的下一个同级标题
                            if header_positions:
                                start_idx = header_positions[0][0]
                                last_idx = header_positions[-1][0]

                                # 找到章节结束位置（下一个同级或更高级的标题）
                                # 获取起始标题的级别（例如 "# 2." 是一级标题）
                                start_line = lines[start_idx].strip()
                                import re
                                start_match = re.match(r'^(#+)\s+(\d+)\.', start_line)
                                if start_match:
                                    start_level_hashes = len(start_match.group(1))
                                    start_number = int(start_match.group(2))

                                    # 从最后一个标题之后开始查找
                                    section_end = len(lines)
                                    for i in range(last_idx + 1, len(lines)):
                                        line_stripped = lines[i].strip()
                                        # 检查是否是同级或更高级的标题
                                        # 匹配 "# 数字." 或 "# 数字.数字." 等格式
                                        match = re.match(r'^(#+)\s+(\d+)(?:\.(\d+))*\.', line_stripped)
                                        if match:
                                            level_hashes = len(match.group(1))
                                            number = int(match.group(2))
                                            sub_number = match.group(3)  # 子编号（如 3.1 中的 1）

                                            # 判断是否是同级或更高级标题：
                                            # 1. 如果 hash 数量相同且没有子编号，且主编号更大 → 同级标题
                                            # 2. 如果 hash 数量更少 → 更高级标题
                                            if level_hashes == start_level_hashes:
                                                # 同级标题：没有子编号，且主编号更大
                                                if sub_number is None and number > start_number:
                                                    section_end = i
                                                    break
                                            elif level_hashes < start_level_hashes:
                                                # 更高级标题
                                                section_end = i
                                                break
                                        # 也检查是否是排除章节（Acknowledgements, References 等）
                                        elif line_stripped.startswith('# ') and not re.match(r'^#+\s+\d+\.', line_stripped):
                                            section_end = i
                                            break
                                else:
                                    # 如果无法解析标题级别，使用简单逻辑：找下一个一级标题
                                    section_end = len(lines)
                                    for i in range(last_idx + 1, len(lines)):
                                        if lines[i].strip().startswith('# ') and not lines[i].strip().startswith('## '):
                                            section_end = i
                                            break

                                # 提取完整内容
                                section_content = '\n'.join(lines[start_idx:section_end])
                                combined_content.append(section_content)
                                self.log(f"       ✓ 提取章节内容: {headers[0][:40]}... 到 {headers[-1][:40]}... → {section_type}")

                            # 合并所有内容
                            if combined_content:
                                new_content = '\n\n'.join(combined_content)
                                if already_exists:
                                    # 如果章节已存在，将新内容添加到前面（因为 LLM 分类的通常是前面的章节）
                                    sections[section_type] = new_content + '\n\n' + sections[section_type]
                                    self.log(f"       ✓ 合并到现有章节: {section_type} (添加 {len(combined_content)} 个片段)")
                                else:
                                    sections[section_type] = new_content
                                    self.log(f"       ✓ 通过分类补充章节: {section_type} (合并 {len(combined_content)} 个片段)")

                        # 重新检查是否还有缺失的关键章节
                        critical_sections = ['Abstract', 'Introduction', 'Methods', 'Results & Discussion', 'Conclusion']
                        missing_critical = [s for s in critical_sections if s not in sections]
                        if missing_critical:
                            self.log(f"    ⓘ 分类后仍缺少: {', '.join(missing_critical)}")
                        else:
                            self.log(f"    ✅ 通过标题分类成功补全所有关键章节！")
                            need_llm = False  # 不需要再用 LLM 提取全文了

                except Exception as e:
                    self.log(f"    ⚠️  标题分类失败: {str(e)}")

            # 如果标题分类后仍有问题，使用 LLM 补充或重新提取
            if need_llm:
                self.log("    ⚠️  质量检查发现问题:")
                for reason in llm_reason:
                    self.log(f"       - {reason}")
                self.log("    🤖 尝试使用 LLM 改进提取结果...")
                self.log("    ⏳ 这可能需要 10-30 秒，请耐心等待...")

                try:
                    # 确定需要提取的章节
                    if not sections:
                        # 如果正则完全失败，提取所有章节
                        missing_sections = None
                        prompt_file = Path(__file__).parent / "prompts" / "section_extraction_prompt.txt"
                        self.log("    📋 正则提取失败，使用 LLM 提取所有章节")
                    else:
                        # 只提取缺失的章节
                        critical_sections = ['Abstract', 'Introduction', 'Methods', 'Results & Discussion', 'Conclusion']
                        missing_sections = [s for s in critical_sections if s not in sections]

                        # 同时包含过短的章节
                        for name, content in sections.items():
                            if len(content.strip()) < 100 and name not in missing_sections:
                                missing_sections.append(name)

                        if missing_sections:
                            prompt_file = Path(__file__).parent / "prompts" / "section_extraction_missing_prompt.txt"
                            self.log(f"    📋 只提取缺失章节: {', '.join(missing_sections)}")
                        else:
                            # 没有缺失章节，不需要调用 LLM
                            self.log("    ⓘ 没有缺失章节，跳过 LLM 调用")
                            need_llm = False

                    if need_llm:
                        # 加载提示词模板
                        if not prompt_file.exists():
                            self.log(f"    ❌ 提示词文件不存在: {prompt_file}")
                        else:
                            self.log("    加载提示词模板...")
                            prompt_template = load_prompt_template(prompt_file)

                            llm_sections = llm.extract_sections(markdown_content, prompt_template, missing_sections)

                            if llm_sections:
                                # 合并策略：优先使用LLM结果，但保留正则表达式的优质结果
                                if not sections:
                                    # 如果正则完全失败，直接使用LLM结果
                                    sections = llm_sections
                                    self.log(f"    ✓ LLM 成功提取 {len(llm_sections)} 个章节")
                                else:
                                    # 智能合并：补充缺失的章节，替换过短的章节
                                    merged_count = 0
                                    for name, llm_content in llm_sections.items():
                                        if name not in sections:
                                            # 补充缺失的章节
                                            sections[name] = llm_content
                                            merged_count += 1
                                            self.log(f"       ✓ 补充章节: {name}")
                                        elif len(sections[name].strip()) < 100 and len(llm_content.strip()) > 100:
                                            # 用LLM的更完整内容替换过短的章节
                                            sections[name] = llm_content
                                            merged_count += 1
                                            self.log(f"       ✓ 改进章节: {name}")

                                    if merged_count > 0:
                                        self.log(f"    ✓ 成功合并 {merged_count} 个章节")
                                    else:
                                        self.log(f"    ⓘ LLM 结果未提供改进")
                            else:
                                self.log("    ⚠️  LLM 提取失败")

                except Exception as e:
                    self.log(f"    ⚠️  LLM 调用出错: {str(e)}")
                    self.log("    ℹ️  将继续使用正则表达式的结果")

            if not sections:
                self.log("    ❌ 所有方法都失败了，无法提取章节")
                self.log("    💡 建议：检查文档格式，或手动提取章节")
                return

            self.log(f"    ✓ 成功识别到 {len(sections)} 个章节")

            # 保存章节
            sections_dir = extract_dir / "Sections"
            saved_files = save_sections(sections, sections_dir)

            self.log(f"    ✓ 保存了 {len(saved_files)} 个章节文件:")
            for file_path in saved_files:
                file_name = Path(file_path).name
                self.log(f"      - {file_name}")

        except Exception as e:
            self.log(f"    ❌ 章节提取失败: {str(e)}")
            import traceback
            self.log(traceback.format_exc())

    def create_word_folder(self, raw_dir: Path, extract_dir: Path, pdf_name: str):
        """创建 Word 文档和 Markdown 图表汇总（按原文顺序排列图片和表格）"""
        try:
            self.log("  - 创建 Word 文档和 Markdown 图表汇总...")

            # 检查是否安装了 python-docx
            try:
                from docx import Document
                from docx.shared import Cm, Pt
                from docx.enum.text import WD_ALIGN_PARAGRAPH
            except ImportError:
                self.log("    ⚠️  需要安装 python-docx 来生成 Word 文档")
                self.log("    ⚠️  运行: pip install python-docx")
                return

            # 读取 Markdown 文件以获取图表标题
            md_file = raw_dir / f"{pdf_name}.md"
            if not md_file.exists():
                self.log("    ⚠️  未找到 Markdown 文件")
                return

            with open(md_file, 'r', encoding='utf-8') as f:
                md_lines = f.readlines()

            # 创建 Word 文档
            doc = Document()

            # 设置页面边距
            sections = doc.sections
            for section in sections:
                section.left_margin = Cm(2.54)
                section.right_margin = Cm(2.54)

            # 创建 Word 文件夹
            word_dir = extract_dir / "Word"
            word_dir.mkdir(exist_ok=True)

            # 获取 Tables 文件夹路径
            tables_dir = extract_dir / "Tables"
            figure_dir = extract_dir / "Figure"

            # 读取图片映射文件
            image_mapping = {}
            mapping_file = figure_dir / "image_mapping.json"
            if mapping_file.exists():
                try:
                    with open(mapping_file, 'r', encoding='utf-8') as f:
                        image_mapping = json.load(f)
                except Exception as e:
                    self.log(f"    ⚠️  读取图片映射文件失败: {str(e)}")

            # 准备 Markdown 图表汇总内容
            md_summary_lines = []
            md_summary_lines.append(f"# {pdf_name} - 图表汇总\n\n")
            md_summary_lines.append("本文档包含从 PDF 中提取的所有图片和表格，按原文顺序排列。\n\n")
            md_summary_lines.append("---\n\n")

            # 遍历 Markdown 文件，按顺序处理图片和表格
            i = 0
            item_count = 0

            while i < len(md_lines):
                line = md_lines[i].strip()

                # 检查是否是表格标题（表格标题在上方）
                if re.match(r'^(Table|表)\s*\d+', line) or (line.startswith('#') and re.search(r'(Table|表)\s*\d+', line)):
                    # 去掉可能存在的#号和空格
                    title = re.sub(r'^#+\s*', '', line)

                    # 添加表格标题（表格标题在表上方）
                    p = doc.add_paragraph(title)
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p_format = p.paragraph_format
                    p_format.space_before = Pt(6)
                    p_format.space_after = Pt(6)

                    # 从标题中提取表格编号
                    table_match = re.search(r'(Table|表)\s*(\d+)', title, re.IGNORECASE)
                    if table_match:
                        table_number = table_match.group(2)

                        # 直接从 Tables 文件夹中查找表格图片
                        table_img = tables_dir / f"Table_{table_number}.jpg"
                        if not table_img.exists():
                            table_img = tables_dir / f"Table_{table_number}.png"

                        if table_img.exists():
                            # 添加表格图片到 Word
                            doc.add_picture(str(table_img), width=Cm(14))
                            doc.add_paragraph()  # 添加空行

                            # 添加到 Markdown 汇总
                            md_summary_lines.append(f"## {title}\n\n")
                            # 使用相对路径引用图片
                            relative_path = f"../Tables/{table_img.name}"
                            md_summary_lines.append(f"![{title}]({relative_path})\n\n")
                            md_summary_lines.append("---\n\n")

                            item_count += 1
                            self.log(f"      添加表格 {table_number}")
                        else:
                            self.log(f"      ⚠️  未找到表格图片: {table_img.name}")

                    i += 1
                    continue

                # 检查是否是图片（先添加图片，再找下一行的标题）
                if line.startswith('!['):
                    # 提取图片路径
                    img_match = re.search(r'\((.*?)\)', line)
                    if img_match:
                        img_path = img_match.group(1)
                        img_full_path = raw_dir / img_path

                        if img_full_path.exists():
                            # 检查这个图片是否是表格图片（已经在表格标题处理过）
                            # 通过检查前面几行是否有表格标题来判断
                            is_table_image = False
                            for j in range(max(0, i-5), i):
                                if re.match(r'^(Table|表)\s*\d+', md_lines[j].strip()):
                                    is_table_image = True
                                    break

                            if not is_table_image:
                                # 添加图片到 Word
                                doc.add_picture(str(img_full_path), width=Cm(14))

                                # 查找下一行的图片标题
                                caption = ""
                                if i + 1 < len(md_lines):
                                    next_line = md_lines[i + 1].strip()
                                    # 匹配图片标题格式
                                    if re.match(r'^(Fig\.|Figure|图|Scheme|图表|示意图|插图)\s*\d+\.?', next_line):
                                        caption = next_line
                                        # 添加图片标题到 Word（图片标题在图下方）
                                        p = doc.add_paragraph(caption)
                                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                        p_format = p.paragraph_format
                                        p_format.space_before = Pt(6)
                                        p_format.space_after = Pt(12)

                                        # 添加到 Markdown 汇总
                                        md_summary_lines.append(f"## {caption}\n\n")

                                        # 使用 Figure 文件夹中的图片路径
                                        if img_path in image_mapping:
                                            # 使用映射后的文件名
                                            new_img_name = image_mapping[img_path]
                                            relative_path = f"../Figure/{new_img_name}"
                                        else:
                                            # 如果没有映射，尝试使用原始路径
                                            relative_path = f"../{img_path}"

                                        md_summary_lines.append(f"![{caption}]({relative_path})\n\n")
                                        md_summary_lines.append("---\n\n")

                                        i += 2  # 跳过标题行
                                        item_count += 1
                                        continue

                                doc.add_paragraph()  # 添加空行
                                item_count += 1

                i += 1

            # 保存 Word 文档
            output_doc = word_dir / f"{pdf_name}_图表.docx"
            doc.save(str(output_doc))

            # 保存 Markdown 图表汇总
            md_summary_file = word_dir / f"{pdf_name}_图表汇总.md"
            with open(md_summary_file, 'w', encoding='utf-8') as f:
                f.writelines(md_summary_lines)

            self.log(f"    ✓ Word 文档创建完成: {output_doc.name}")
            self.log(f"    ✓ Markdown 图表汇总创建完成: {md_summary_file.name}")
            self.log(f"    ✓ 包含 {item_count} 个图表")

        except Exception as e:
            self.log(f"    ❌ Word 文档创建失败: {str(e)}")
            import traceback
            self.log(traceback.format_exc())


def main():
    """主函数"""
    root = tk.Tk()
    BatchPDFProcessorGUI(root)
    root.mainloop()


if __name__ == "__main__":
    main()
